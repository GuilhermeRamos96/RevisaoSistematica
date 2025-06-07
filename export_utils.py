# -*- coding: utf-8 -*-
"""
Export Utilities for PRISMA 2020 Tool
Handles document generation in DOCX, PDF, and JSON formats
"""

import json
import os
from typing import Dict, List, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available. DOCX export disabled.")
    DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    print("Warning: reportlab not available. PDF export disabled.")
    PDF_AVAILABLE = False

class ExportUtils:
    """
    Comprehensive export utilities for PRISMA 2020 systematic reviews
    Supports DOCX, PDF, and JSON export formats
    """
    
    def __init__(self):
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
    def export_to_docx(self, content_data: Dict[str, Any]) -> Optional[str]:
        """
        Export systematic review to DOCX format
        
        Args:
            content_data: Dictionary containing sections, references, and citations
            
        Returns:
            Path to generated DOCX file or None if failed
        """
        if not DOCX_AVAILABLE:
            print("DOCX export not available - python-docx not installed")
            return None
            
        try:
            doc = Document()
            
            # Document title
            title = doc.add_heading('Systematic Review - PRISMA 2020', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add generation date
            date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_page_break()
            
            # Add sections
            sections = content_data.get('sections', {})
            for section_name, section_content in sections.items():
                if section_content.strip():  # Only add non-empty sections
                    # Section heading
                    doc.add_heading(section_name, level=1)
                    
                    # Section content
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    
                    doc.add_paragraph()  # Add spacing
            
            # Add references section
            references = content_data.get('references', [])
            if references:
                doc.add_page_break()
                doc.add_heading('References', level=1)
                
                for ref in references:
                    ref_para = doc.add_paragraph()
                    ref_para.add_run(f"{ref['id']}. ").bold = True
                    ref_para.add_run(ref['formatted'])
                    ref_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Save document
            filename = f'prisma_review_{self.timestamp}.docx'
            doc.save(filename)
            print(f"DOCX export saved as: {filename}")
            return filename
            
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def export_to_pdf(self, content_data: Dict[str, Any]) -> Optional[str]:
        """
        Export systematic review to PDF format
        
        Args:
            content_data: Dictionary containing sections, references, and citations
            
        Returns:
            Path to generated PDF file or None if failed
        """
        if not PDF_AVAILABLE:
            print("PDF export not available - reportlab not installed")
            return None
            
        try:
            filename = f'prisma_review_{self.timestamp}.pdf'
            doc = SimpleDocTemplate(filename, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=20
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                alignment=TA_JUSTIFY
            )
            
            # Document title
            story.append(Paragraph("Systematic Review - PRISMA 2020", title_style))
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
            story.append(Spacer(1, 0.5*inch))
            
            # Add sections
            sections = content_data.get('sections', {})
            for section_name, section_content in sections.items():
                if section_content.strip():  # Only add non-empty sections
                    # Section heading
                    story.append(Paragraph(section_name, heading_style))
                    
                    # Section content
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            # Handle citations in text
                            formatted_text = self._format_text_for_pdf(para_text.strip())
                            story.append(Paragraph(formatted_text, body_style))
                    
                    story.append(Spacer(1, 0.2*inch))
            
            # Add references section
            references = content_data.get('references', [])
            if references:
                story.append(PageBreak())
                story.append(Paragraph("References", heading_style))
                
                for ref in references:
                    ref_text = f"<b>{ref['id']}.</b> {ref['formatted']}"
                    story.append(Paragraph(ref_text, body_style))
            
            # Build PDF
            doc.build(story)
            print(f"PDF export saved as: {filename}")
            return filename
            
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def export_to_json(self, content_data: Dict[str, Any]) -> Optional[str]:
        """
        Export systematic review to JSON format
        
        Args:
            content_data: Dictionary containing sections, references, and citations
            
        Returns:
            JSON string or None if failed
        """
        try:
            export_data = {
                'metadata': {
                    'export_timestamp': datetime.now().isoformat(),
                    'format_version': '1.0',
                    'tool': 'PRISMA 2020 Systematic Review Tool'
                },
                'sections': content_data.get('sections', {}),
                'references': content_data.get('references', []),
                'citations': content_data.get('citations', {}),
                'statistics': self._calculate_statistics(content_data)
            }
            
            json_string = json.dumps(export_data, indent=2, ensure_ascii=False)
            print("JSON export generated successfully")
            return json_string
            
        except Exception as e:
            print(f"Error exporting to JSON: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def import_from_json(self, json_data: str) -> Optional[Dict[str, Any]]:
        """
        Import systematic review from JSON format
        
        Args:
            json_data: JSON string containing review data
            
        Returns:
            Dictionary with imported data or None if failed
        """
        try:
            data = json.loads(json_data)
            
            # Validate format
            if 'sections' not in data or 'references' not in data:
                print("Invalid JSON format - missing required fields")
                return None
            
            imported_data = {
                'sections': data.get('sections', {}),
                'references': data.get('references', []),
                'citations': data.get('citations', {})
            }
            
            print("JSON import successful")
            return imported_data
            
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON: {e}")
            return None
        except Exception as e:
            print(f"Error importing from JSON: {e}")
            return None
    
    def _format_text_for_pdf(self, text: str) -> str:
        """
        Format text for PDF export, handling special characters and citations
        
        Args:
            text: Raw text to format
            
        Returns:
            Formatted text suitable for PDF
        """
        # Escape special characters for reportlab
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        # Format citations (simple approach - could be enhanced)
        import re
        citation_pattern = r'\[(\d+)\]'
        text = re.sub(citation_pattern, r'<sup>[\1]</sup>', text)
        
        return text
    
    def _calculate_statistics(self, content_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Calculate statistics about the review content
        
        Args:
            content_data: Dictionary containing review data
            
        Returns:
            Dictionary with calculated statistics
        """
        sections = content_data.get('sections', {})
        references = content_data.get('references', [])
        citations = content_data.get('citations', {})
        
        # Section statistics
        completed_sections = sum(1 for content in sections.values() if content.strip())
        total_words = sum(len(content.split()) for content in sections.values() if content.strip())
        
        # Reference statistics
        total_references = len(references)
        cited_references = set()
        for ref_list in citations.values():
            cited_references.update(ref_list)
        
        return {
            'sections': {
                'total': len(sections),
                'completed': completed_sections,
                'completion_rate': completed_sections / len(sections) if sections else 0,
                'total_words': total_words
            },
            'references': {
                'total': total_references,
                'cited': len(cited_references),
                'uncited': total_references - len(cited_references)
            },
            'citations': {
                'total_citations': sum(len(ref_list) for ref_list in citations.values()),
                'sections_with_citations': len([s for s in citations.values() if s])
            }
        }
    
    def generate_checklist_docx(self) -> Optional[str]:
        """
        Generate PRISMA 2020 checklist document
        
        Returns:
            Path to generated checklist DOCX file or None if failed
        """
        if not DOCX_AVAILABLE:
            print("DOCX export not available - python-docx not installed")
            return None
            
        try:
            doc = Document()
            
            # Document title
            title = doc.add_heading('PRISMA 2020 Checklist', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add introduction
            intro = doc.add_paragraph(
                "This checklist has been adapted from the PRISMA 2020 statement which can be found at: "
                "http://www.prisma-statement.org/. When completing this checklist consider the PRISMA 2020 "
                "Explanation and Elaboration paper for important clarification on these items."
            )
            intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            doc.add_paragraph()
            
            # Define PRISMA 2020 checklist items
            checklist_items = [
                {
                    "section": "TITLE",
                    "items": [
                        {"item": "1", "description": "Identify the report as a systematic review."}
                    ]
                },
                {
                    "section": "ABSTRACT",
                    "items": [
                        {"item": "2", "description": "See the PRISMA 2020 for Abstracts checklist."}
                    ]
                },
                {
                    "section": "INTRODUCTION",
                    "items": [
                        {"item": "3", "description": "Describe the rationale for the review in the context of existing knowledge."},
                        {"item": "4", "description": "Provide an explicit statement of the objective(s) or question(s) the review addresses."}
                    ]
                },
                {
                    "section": "METHODS",
                    "items": [
                        {"item": "5", "description": "Indicate whether a review protocol exists; provide registration information including registration number."},
                        {"item": "6", "description": "Specify characteristics of the studies and other sources that were eligible for inclusion."},
                        {"item": "7", "description": "Describe all information sources and date last searched or consulted."},
                        {"item": "8", "description": "Present the full electronic search strategy for at least 1 database."},
                        {"item": "9", "description": "State the process for selecting studies."},
                        {"item": "10", "description": "Describe the method of data extraction from reports."},
                        {"item": "11", "description": "List and define all variables for which data were sought."},
                        {"item": "12", "description": "Describe methods used for assessing risk of bias of individual studies."},
                        {"item": "13", "description": "State the principal summary measures."},
                        {"item": "14", "description": "Describe the methods of handling data and combining results of studies."},
                        {"item": "15", "description": "Specify any assessment of risk of bias that may affect the cumulative evidence."},
                        {"item": "16", "description": "Describe methods used to assess the certainty of the evidence."}
                    ]
                },
                {
                    "section": "RESULTS",
                    "items": [
                        {"item": "17", "description": "Give numbers of studies screened, assessed for eligibility, and included in the review, with reasons for exclusions at each stage, ideally with a flow diagram."},
                        {"item": "18", "description": "Present characteristics of included studies and provide the citations."},
                        {"item": "19", "description": "Present assessments of risk of bias for each included study."},
                        {"item": "20", "description": "For all outcomes, and for each study, present simple summary data and effect estimates and confidence intervals."},
                        {"item": "21", "description": "Present results of each meta-analysis done, including confidence intervals and measures of consistency."},
                        {"item": "22", "description": "Present results of any assessment of risk of bias across studies."},
                        {"item": "23", "description": "Present results of additional analyses, if done."},
                        {"item": "24", "description": "Present assessments of certainty of evidence for each outcome."}
                    ]
                },
                {
                    "section": "DISCUSSION",
                    "items": [
                        {"item": "25", "description": "Provide a general interpretation of the results in the context of other evidence."},
                        {"item": "26", "description": "Discuss limitations at study and outcome level, and at review level."},
                        {"item": "27", "description": "Provide a general interpretation of the results in the context of other evidence, and implications for future research."}
                    ]
                },
                {
                    "section": "OTHER INFORMATION",
                    "items": [
                        {"item": "28", "description": "Describe sources of funding for the systematic review and other support."},
                        {"item": "29", "description": "Declare any conflicts of interest of review authors."},
                        {"item": "30", "description": "Report how the review was funded and the role of the funders."}
                    ]
                }
            ]
            
            # Create table for checklist
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Item #'
            header_cells[1].text = 'PRISMA 2020 Item'
            header_cells[2].text = 'Location in manuscript'
            header_cells[3].text = 'Location in your review'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add checklist items
            for section_data in checklist_items:
                # Add section header
                section_row = table.add_row()
                section_cells = section_row.cells
                section_cells[0].text = ""
                section_cells[1].text = section_data["section"]
                section_cells[2].text = ""
                section_cells[3].text = ""
                
                # Make section header bold
                for paragraph in section_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                
                # Add items
                for item_data in section_data["items"]:
                    item_row = table.add_row()
                    item_cells = item_row.cells
                    item_cells[0].text = item_data["item"]
                    item_cells[1].text = item_data["description"]
                    item_cells[2].text = "Page: ___"
                    item_cells[3].text = ""
            
            # Save document
            filename = f'prisma_checklist_{self.timestamp}.docx'
            doc.save(filename)
            print(f"PRISMA checklist saved as: {filename}")
            return filename
            
        except Exception as e:
            print(f"Error generating PRISMA checklist: {e}")
            import traceback
            traceback.print_exc()
            return None
